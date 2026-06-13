from claude_client import (
    generate_proposal_narrative,
    generate_directives,
    generate_gonogo_reasoning
)
from rag_engine import get_compliance_stats


def build_proposal(
    requirements: dict,
    compliance_items: list[dict],
    scoring_result: dict
) -> dict:
    """
    Orchestrates all Claude calls to produce the final proposal package.

    Args:
        requirements:     output from parser.parse_rfp()["requirements"]
        compliance_items: output from rag_engine.run_compliance_check()
        scoring_result:   output from scoring.calculate_win_probability()

    Returns complete dict matching the frontend/Supabase schema:
    {
        "win_probability":  int,
        "budget_score":     int,
        "capability_score": int,
        "decision":         str,
        "compliance":       [...],
        "directives":       [...],
        "proposal_narrative": str,
        "factor_breakdown": {           ← included for frontend gauge bars
            "compliance_pts": float,
            "gaps_pts":       float,
            "sector_pts":     float,
            "budget_pts":     float
        }
    }
    """
    # Extract values from scoring result
    win_probability   = scoring_result["win_probability"]
    budget_score      = scoring_result["budget_score"]
    capability_score  = scoring_result["capability_score"]
    decision          = scoring_result["decision"]
    factor_breakdown  = scoring_result.get("factor_breakdown", {})

    # Get compliance stats for reasoning
    stats          = get_compliance_stats(compliance_items)
    compliance_pct = stats["compliance_pct"]
    gaps_count     = stats["gaps_count"]
    sector         = requirements.get("sector", "Other")

    # ── Claude Call 1: Proposal Narrative ──────────────────────
    proposal_narrative = generate_proposal_narrative(
        requirements=requirements.get("mandatory_requirements", []),
        compliance_items=compliance_items,
        sector=sector,
        win_probability=win_probability,
        decision=decision
    )

    # ── Claude Call 2: Strategic Directives ────────────────────
    directives = generate_directives(
        requirements=requirements.get("mandatory_requirements", []),
        compliance_items=compliance_items,
        decision=decision
    )

    # ── Claude Call 3: GO/NO-GO Reasoning ──────────────────────
    reasoning = generate_gonogo_reasoning(
        win_probability=win_probability,
        decision=decision,
        sector=sector,
        compliance_pct=compliance_pct,
        gaps_count=gaps_count
    )

    full_narrative = proposal_narrative

    return {
        "win_probability":   win_probability,
        "budget_score":      budget_score,
        "capability_score":  capability_score,
        "decision":          decision,
        "compliance":        compliance_items,
        "directives":        directives,
        "proposal_narrative": full_narrative,
        # ── NOT stored in Supabase — only returned in API response ──
        "factor_breakdown":  factor_breakdown
    }


def export_proposal_to_docx(proposal: dict, output_path: str = "proposal_output.docx") -> str:
    """
    Export proposal narrative + compliance table to a Word document.
    Returns the output file path.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    doc = Document()

    # ── Title ──────────────────────────────────────────────────
    title = doc.add_heading("TEKROWE — Proposal Response", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Decision Banner ────────────────────────────────────────
    banner = doc.add_paragraph()
    banner_run = banner.add_run(
        f"Decision: {proposal['decision']}  |  "
        f"Win Probability: {proposal['win_probability']}%  |  "
        f"Capability Score: {proposal['capability_score']}  |  "
        f"Budget Score: {proposal['budget_score']}"
    )
    banner_run.bold = True
    banner_run.font.size = Pt(11)
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── Proposal Narrative ─────────────────────────────────────
    doc.add_heading("Proposal Narrative", level=1)
    clean_narrative = re.sub(r"\*\*(.*?)\*\*", r"\1", proposal["proposal_narrative"])
    for para_text in clean_narrative.split("\n\n"):
        if para_text.strip():
            if para_text.startswith("---"):
                doc.add_paragraph("─" * 60)
            else:
                doc.add_paragraph(para_text.strip())
    doc.add_paragraph()

    # ── Strategic Directives ───────────────────────────────────
    doc.add_heading("Strategic Directives", level=1)
    for directive_html in proposal["directives"]:
        clean = re.sub(r"<[^>]+>", "", directive_html).replace("<br>", "\n")
        doc.add_paragraph(clean.strip(), style="List Bullet")
    doc.add_paragraph()

    # ── Compliance Matrix ──────────────────────────────────────
    doc.add_heading("Compliance Matrix", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["ID", "Requirement", "Matched Capability", "Status"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    for item in proposal["compliance"]:
        row_cells = table.add_row().cells
        row_cells[0].text = item.get("id", "")
        row_cells[1].text = item.get("req", "")
        row_cells[2].text = item.get("match", "N/A")
        status = item.get("status", "")
        row_cells[3].text = status
        run = row_cells[3].paragraphs[0].runs
        if run:
            if status == "COMPLIANT":
                run[0].font.color.rgb = RGBColor(0, 128, 0)
            else:
                run[0].font.color.rgb = RGBColor(200, 0, 0)

    doc.save(output_path)
    return output_path
